"""生成第一次向程教授汇报使用的 AI_PIC_Demo 简洁介绍说明书。"""

from __future__ import annotations

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    print("缺少 python-docx，请先安装：pip install python-docx", file=sys.stderr)
    raise SystemExit(1)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "docs" / "程教授汇报_Demo介绍说明书.docx"
IMAGE_ROOT_NAMES = ("outputs", "output", "reports", "report", "results", "figures", "docs")
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg"}


def set_run_font(run, name: str = "宋体", size: float = 12, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_begin, instr, fld_end))
    set_run_font(run, size=9)


def add_title(document: Document, version: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("基于 AI 大模型的光子芯片设计平台 Demo 介绍说明书")
    set_run_font(run, name="黑体", size=16, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(3)
    run = subtitle.add_run("——第一次阶段性汇报材料")
    set_run_font(run, name="宋体", size=12, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run(f"项目：AI_PIC_Demo　　当前版本：{version}")
    set_run_font(run, size=10.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, name="黑体", size=14 if level == 1 else 12, bold=True)


def add_body(document: Document, text: str, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.first_line_indent = Pt(24) if indent else Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_numbered_items(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.left_indent = Pt(20)
        paragraph.paragraph_format.first_line_indent = Pt(-20)
        run = paragraph.add_run(f"{index}. {item}")
        set_run_font(run, size=11.5)


def add_compact_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9D9D9")
        table.rows[0].cells[index].width = Cm(widths[index])
    prevent_row_split(table.rows[0])
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value)
            row.cells[index].width = Cm(widths[index])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def iter_candidate_images() -> list[Path]:
    candidates: list[Path] = []
    for name in IMAGE_ROOT_NAMES:
        root = ROOT / name
        if root.exists():
            candidates.extend(
                path
                for path in root.rglob("*")
                if path.is_file() and path.suffix.lower() in IMAGE_SUFFIXES
            )
    return candidates


def choose_image(candidates: list[Path], preferred_names: tuple[str, ...]) -> Path | None:
    ranked: list[tuple[int, float, Path]] = []
    for path in candidates:
        lower_name = path.name.lower()
        for priority, preferred in enumerate(preferred_names):
            if preferred in lower_name:
                exact_bonus = 100 if path.stem.lower() == preferred else 0
                ranked.append((exact_bonus + len(preferred_names) - priority, path.stat().st_mtime, path))
                break
    return max(ranked, default=(0, 0.0, None), key=lambda item: (item[0], item[1]))[2]


def select_demo_images() -> list[tuple[str, Path]]:
    candidates = iter_candidate_images()
    specifications = [
        ("图1  MMI 长度扫描结果", ("length_sweep", "length", "sweep")),
        ("图2  MMI 宽度-长度二维优化结果", ("width_length_heatmap", "heatmap", "scan")),
        ("图3  电磁场传播结果示意", ("field_propagation_enhanced", "field_propagation", "propagation", "field")),
        ("图4  1×2 MMI 版图预览", ("layout_preview", "gds", "layout", "mmi")),
    ]
    selected: list[tuple[str, Path]] = []
    used: set[Path] = set()
    for caption, names in specifications:
        image = choose_image(candidates, names)
        if image is not None and image not in used:
            selected.append((caption, image))
            used.add(image)
    return selected


def add_figure_grid(document: Document, images: list[tuple[str, Path]]) -> None:
    if not images:
        add_body(document, "当前未检索到可插入的项目结果图片，本节保留文字说明。")
        return
    table = document.add_table(rows=(len(images) + 1) // 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (caption, path) in enumerate(images):
        cell = table.cell(index // 2, index % 2)
        cell.width = Cm(7.4)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        try:
            run.add_picture(str(path), width=Cm(7.0))
        except Exception as exc:
            set_run_font(run, size=10)
            run.text = f"图片读取失败：{path.name}（{exc}）"
        caption_paragraph = cell.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_before = Pt(2)
        caption_paragraph.paragraph_format.space_after = Pt(3)
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(caption_run, size=10)
    if len(images) % 2:
        table.cell(len(images) // 2, 1).text = ""


def read_version() -> str:
    version_path = ROOT / "VERSION"
    return version_path.read_text(encoding="utf-8").strip() if version_path.exists() else "当前开发版本"


def load_latest_result_summary() -> str:
    run_dirs = sorted((ROOT / "outputs").glob("run_*"), key=lambda path: path.stat().st_mtime, reverse=True)
    for run_dir in run_dirs:
        result_path = run_dir / "optimization_result.json"
        if not result_path.exists():
            continue
        try:
            result = json.loads(result_path.read_text(encoding="utf-8"))
            width = float(result["best_width_um"])
            length = float(result["best_length_um"])
            return f"最近完整结果示例给出的 MMI 优选宽度约为 {width:.3f} μm、长度约为 {length:.3f} μm；该数值来自当前简化优化模型，仅用于 Demo 原型展示。"
        except (KeyError, TypeError, ValueError, json.JSONDecodeError):
            continue
    return "当前版本能够输出参数扫描结果；具体数值应以每次运行目录中的结构化结果为准。"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.footer_distance = Cm(1.2)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def build_document() -> tuple[Path, list[tuple[str, Path]]]:
    version = read_version()
    images = select_demo_images()
    document = Document()
    configure_document(document)
    add_title(document, version)

    add_heading(document, "一、项目背景与目标")
    add_body(
        document,
        "光子芯片设计通常涉及需求理解、器件建模、电磁场求解、参数扫描优化、版图生成和结果报告等多个环节。本项目尝试探索 AI 大模型与 AI Agent 在光子芯片自动化设计流程中的作用，当前选择 SOI 平台上的 1×2 MMI 分束器作为示例器件，验证从自然语言需求到初步设计结果输出的自动化闭环。",
    )
    add_body(
        document,
        "当前 Demo 的重点是打通流程与验证模块接口，并非替代商业级光子 EDA。自然语言解析目前为关键词与正则表达式驱动的规则版原型，已预留后续替换为大模型结构化输出和 Agent 工具调用的接口。",
    )

    add_heading(document, "二、当前 Demo 实现的核心流程")
    add_compact_table(
        document,
        ["阶段", "当前实现", "主要输出"],
        [
            ["需求输入与解析", "自然语言输入，解析器件、平台、波长和目标分光比", "design_spec.json"],
            ["物理参数与建模", "SOI 材料配置、波导截面与标量模式求解、MMI 初始建模", "physical_params.json、mode_result.json"],
            ["扫描与传播验证", "长度扫描、宽度-长度二维扫描、二维标量 BPM 传播", "扫描图、传播场、结构化结果"],
            ["版图与交付", "选择较优参数，生成 GDS、图表、Markdown 报告和结果包", "GDS、PNG、report.md、ZIP"],
        ],
        [3.2, 7.5, 4.6],
    )
    add_body(
        document,
        "总体链路为：自然语言需求输入 → 结构化参数解析 → SOI 材料/器件参数配置 → MMI 器件建模 → 电磁场传播或简化求解 → 长度扫描与二维参数扫描 → 较优参数选择 → GDS 版图生成 → 图表与报告输出。",
        indent=False,
    )

    document.add_page_break()
    add_heading(document, "三、当前已完成的主要功能")
    add_compact_table(
        document,
        ["模块", "已实现内容", "后续改进方向"],
        [
            ["自然语言需求解析", "可从类似“1×2 MMI、1550 nm、SOI、50/50 分束”的描述中提取结构化参数；TE/TM 字段尚未独立解析", "接入大模型结构化输出，补充模式、容差和约束解析"],
            ["SOI 平台与材料", "配置 Si/SiO₂ 等材料折射率及波导几何参数，并生成折射率截面", "加入色散、温度和真实工艺 LayerStack"],
            ["1×2 MMI 建模", "包含输入波导、多模区和双输出结构参数，使用简化自成像模型给出初值", "扩展制造偏差与更多结构自由度"],
            ["传播与模式求解", "实现二维标量有限差分模式求解与二维标量 BPM 传播，可输出模式场和传播场", "与全矢量 FDE/FDTD/FEM/EME 对比验证"],
            ["参数扫描与优化", "实现 MMI 长度扫描、宽度-长度二维联合扫描和波长趋势分析", "引入贝叶斯、遗传、伴随优化或逆向设计"],
            ["端口功率分析", "已实现窗口积分、窗口敏感性和简化 Gaussian 端口模式重叠", "用真实本征模式替代 Gaussian，并形成宽带 S 参数趋势"],
            ["GDS 版图生成", "通过 gdsfactory generic PDK 生成 1×2 MMI GDS 与结构预览", "接入真实 PDK、DRC 与工艺约束"],
            ["报告与网页", "Streamlit 展示结果；自动输出 Markdown、图表、日志和 ZIP 结果包", "扩展自动 Word/PDF 与 Agent 汇报生成"],
        ],
        [3.0, 7.7, 4.6],
    )
    add_body(document, load_latest_result_summary())

    document.add_page_break()
    add_heading(document, "四、目前 Demo 的展示效果")
    add_body(
        document,
        "以下图片由脚本从项目结果目录自动检索并插入，优先采用最近一次完整运行生成的扫描、传播和版图结果。",
    )
    add_figure_grid(document, images)

    document.add_page_break()
    add_heading(document, "五、当前 Demo 的技术特点")
    add_numbered_items(
        document,
        [
            "初步形成从自然语言需求到设计结果、GDS 和报告输出的自动化链路。",
            "以 1×2 MMI 为最小案例，将光子器件设计流程拆分为可单独替换和调用的模块。",
            "模式求解、快速替代模型、BPM 传播和端口模式重叠构成分层验证链路。",
            "代码结构便于后续接入高精度求解器、真实 PDK、优化算法及 AI Agent 工作流。",
            "当前定位是原型验证，重点在流程打通与接口验证，而不是最终工程精度。",
        ],
    )

    add_heading(document, "六、当前局限与需要改进的问题")
    add_numbered_items(
        document,
        [
            "电磁场求解仍为二维标量有限差分和 BPM 近似，尚非商业级全矢量 FDTD/FEM/EME；结果需与高精度方法或实验对比。",
            "器件类型目前主要集中在 1×2 MMI，波导、光栅耦合器、微环、调制器和探测器等尚未形成完整器件库。",
            "AI 作用目前主要体现为规则版需求解析接口和流程组织；真正的大模型 API、自主 Agent 规划、结果检查与迭代闭环尚未完成。",
            "参数优化仍以规则网格扫描和轻量 surrogate model 为主，尚未接入贝叶斯优化、遗传算法、伴随优化或逆向设计。",
            "GDS 当前基于 gdsfactory generic PDK，尚未结合真实代工 PDK、DRC 和完整工艺约束。",
        ],
    )

    document.add_page_break()
    add_heading(document, "七、下一步工作计划")
    add_numbered_items(
        document,
        [
            "完善端口模式重叠积分：在现有 Gaussian 近似基础上接入真实端口本征模式，并开展多波长功率与分束比分析。",
            "将当前求解模块与标准仿真工具或更高精度数值方法进行结果对比，建立误差和适用范围说明。",
            "增加 AI Agent 工作流，使其能够自动拆解任务、调用仿真、读取结果、检查异常并生成下一轮优化建议。",
            "扩展器件库，从 1×2 MMI 逐步增加基础波导、弯曲波导、光栅耦合器和微环谐振器。",
            "继续完善 GitHub 版本管理、版本说明、代表性示例和汇报材料，便于导师与课题组成员查看进展。",
        ],
    )

    add_heading(document, "八、总结")
    add_body(
        document,
        "当前 Demo 已经初步打通“自然语言需求输入—器件参数解析—建模仿真—参数优化—GDS 版图—报告输出”的原型流程，证明该方向具备继续推进的可行性。后续重点是提高物理仿真可信度、增强 AI Agent 自动化能力，并逐步扩展到更多光子器件和更真实的工艺约束。",
    )
    add_body(
        document,
        "阶段性判断：本项目已具备第一次向导师展示整体路线、模块实现和实际输出的条件；汇报时应主动强调“原型验证”和“分层近似模型”的边界。",
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH, images


def main() -> None:
    output_path, images = build_document()
    print(f"文档已生成：{output_path}")
    if images:
        print("使用的项目图片：")
        for caption, path in images:
            print(f"- {caption}：{path.relative_to(ROOT)}")
    else:
        print("未找到适合插入的项目图片，文档已使用纯文字说明。")


if __name__ == "__main__":
    main()
